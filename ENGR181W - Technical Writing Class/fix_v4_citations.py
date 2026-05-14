"""
Apply citation accuracy fixes to Ethics_Impact_Essay_v4.docx.

Changes:
  [1] Bird citation - rewrite city/climate claim to match what Bird actually supports
  [5] IEC 62133-2 - clarify framing so 70 us is attributed to BQ76930 spec
  [9] CPSC - rewrite to match what Dec 2022 letter actually said
  [2] URL - point to actual lithium-ion safety page
  [6] URL - drop trailing .html
  [12] TI datasheet - correct doc number (SLUSBK2I), correct title, correct date
"""
from docx import Document

REPLACEMENTS = [
    # Body: Bird citation - the original cities (LA, Amsterdam, Shenzhen, Nairobi)
    # and climate-specifics aren't in Bird's report. Bird does document global
    # multi-city deployment, so keep [1] but loosen the claim to what's supported.
    (
        "Shared scooter and e-bike fleets operate across Los Angeles, Amsterdam, "
        "Shenzhen, and Nairobi, cities with sharply different climate profiles, "
        "from tropical humidity to sub-freezing winters [1].",
        "Shared scooter and e-bike fleets now operate in cities across multiple "
        "continents [1], spanning climate profiles from tropical humidity to "
        "sub-freezing winters."
    ),

    # Body: 70 us / IEC 62133-2 framing. The 70 us figure is from the BQ76930
    # spec, not from IEC 62133-2. Reword so [5] is cited for the safety
    # framework it actually defines (short-circuit and overcharge tests).
    (
        "short-circuit protection trips within 70 µs with no software "
        "involvement, consistent with IEC 62133-2 requirements for portable "
        "lithium cells [5].",
        "short-circuit protection trips within 70 µs with no software "
        "involvement, in line with the short-circuit and overcharge safety "
        "framework IEC 62133-2 defines for portable lithium cells [5]."
    ),

    # Body: CPSC framing. The Dec 2022 letter urged voluntary UL compliance;
    # CPSC moved toward mandatory standards later. Rewrite to match what [9]
    # actually says.
    (
        "The CPSC is developing mandatory US safety standards for e-mobility "
        "devices following a series of high-profile fires [9].",
        "The CPSC has called on manufacturers, importers, and retailers of "
        "lithium-ion batteries and e-mobility devices to certify compliance "
        "with consensus UL safety standards following a series of high-profile "
        "fires [9], and has since moved toward mandatory US safety standards."
    ),

    # Reference [2]: URL too generic
    (
        "[2] New York City Fire Department, “FDNY Public Safety Bulletin: "
        "Lithium-Ion Battery Fires,” New York, NY, USA, 2023. [Online]. "
        "Available: https://www.nyc.gov/site/fdny/",
        "[2] New York City Fire Department, “Lithium-Ion Battery Safety,” "
        "New York, NY, USA, 2023. [Online]. Available: "
        "https://www.nyc.gov/site/fdny/codes/reference/lithium-ion-battery-safety.page"
    ),

    # Reference [6]: drop stray .html
    (
        "https://www.ieee.org/about/corporate/governance/p7-8.html",
        "https://www.ieee.org/about/corporate/governance/p7-8"
    ),

    # Reference [12]: correct datasheet number, title, and date
    (
        "[12] Texas Instruments, “BQ76930 3-Series to 15-Series Cell "
        "Battery Monitor and Protector,” Datasheet SLUSBK3, Dallas, TX, "
        "USA, 2018. [Online]. Available: https://www.ti.com/product/BQ76930",
        "[12] Texas Instruments, “BQ769x0 3-Series to 15-Series Cell "
        "Battery Monitor Family for Li-Ion and Phosphate Applications,” "
        "Datasheet SLUSBK2I, Dallas, TX, USA, Mar. 2022. [Online]. Available: "
        "https://www.ti.com/product/BQ76930"
    ),
]


def replace_in_paragraph(p, old, new):
    """Run-aware text replace. Preserves first-run formatting."""
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    new_text = full.replace(old, new)
    if not p.runs:
        return False
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""
    return True


def walk_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def main():
    src = "Ethics_Impact_Essay_v4.docx"
    dst = "Ethics_Impact_Essay_v4.docx"
    doc = Document(src)
    applied = {old: 0 for old, _ in REPLACEMENTS}
    for p in walk_paragraphs(doc):
        for old, new in REPLACEMENTS:
            if replace_in_paragraph(p, old, new):
                applied[old] += 1
    try:
        doc.save(dst)
    except PermissionError:
        dst = "Ethics_Impact_Essay_v4_citations_fixed.docx"
        doc.save(dst)
        print("(v4 was locked — wrote to", dst, "instead)")
    print("Saved", dst)
    for i, (old, _) in enumerate(REPLACEMENTS, 1):
        head = old[:60].replace("\n", " ")
        print(f"  fix {i}: {applied[old]} replacement(s) - {head}...")


if __name__ == "__main__":
    main()
