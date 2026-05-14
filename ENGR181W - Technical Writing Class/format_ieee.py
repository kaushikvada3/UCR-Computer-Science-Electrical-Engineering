"""Rebuild Ethics_Impact_Essay.docx in IEEEtran journal style."""
import docx as docxold
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = 'Ethics_Impact_Essay.docx'

# 1. Pull text out of the current doc, classify each paragraph.
src = docxold.Document(SRC)
src_paras = [p.text for p in src.paragraphs if p.text.strip()]

# Find indices for known anchors
def find_idx(prefix):
    for i, t in enumerate(src_paras):
        if t.strip().startswith(prefix):
            return i
    raise ValueError(prefix)

i_abs   = find_idx('Abstract')
i_kw    = find_idx('Keywords')
i_intro = find_idx('1. Introduction')
i_211   = find_idx('2.1.1')
i_212   = find_idx('2.1.2')
i_221   = find_idx('2.2.1')
i_eth   = find_idx('Ethical Implications of Commercialization')
i_les   = find_idx('Professional and Ethical Responsibility')
i_imp   = find_idx('Global, Economic, Environmental')
i_con   = find_idx('Conclusion')
i_ref   = find_idx('References')

abstract_text = src_paras[i_abs + 1]
keywords_text = src_paras[i_kw + 1]

def slice_body(start, end):
    # Take paragraphs strictly between start and end, skipping the heading at start.
    return src_paras[start + 1:end]

intro_body = slice_body(i_intro, i_211)
sub_a_body = slice_body(i_211,   i_212)
sub_b_body = slice_body(i_212,   i_221)
sub_c_body = slice_body(i_221,   i_eth)
eth_body   = slice_body(i_eth,   i_les)
les_body   = slice_body(i_les,   i_imp)
imp_body   = slice_body(i_imp,   i_con)
con_body   = slice_body(i_con,   i_ref)
references = src_paras[i_ref + 1:]

title_lines = [
    "Understanding of Professional and Ethical Responsibility "
    "and Global, Economic, Environmental, and Societal Impact",
]
subtitle = "Modular Battery Management and Communication Platform"
authors  = "Kaushik Vada, Connor Stewart, Nick Poon, Christian Kim"
affil    = "UC Riverside, EE 175AB, 2026"

# 2. Build new doc with IEEEtran-style formatting.
doc = Document()

# Default font
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(10)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
    rfonts.set(qn(attr), 'Times New Roman')

# Page margins (IEEE journal: ~0.75 top, 1 bottom, 0.625 sides on Letter)
sec = doc.sections[0]
sec.top_margin    = Inches(0.75)
sec.bottom_margin = Inches(1.0)
sec.left_margin   = Inches(0.625)
sec.right_margin  = Inches(0.625)

# ----- Title block (full width) -----
def center_para(text, size, bold=False, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

center_para(title_lines[0], 20, space_after=6)
center_para(subtitle, 14, space_after=6)
center_para(authors, 11, italic=True, space_after=2)
center_para(affil, 10, space_after=10)

# ----- Continuous section break -> two columns -----
new_section = doc.add_section(WD_SECTION.CONTINUOUS)
new_section.top_margin    = Inches(0.75)
new_section.bottom_margin = Inches(1.0)
new_section.left_margin   = Inches(0.625)
new_section.right_margin  = Inches(0.625)
sectPr = new_section._sectPr
cols = sectPr.find(qn('w:cols'))
if cols is None:
    cols = OxmlElement('w:cols')
    sectPr.append(cols)
cols.set(qn('w:num'), '2')
cols.set(qn('w:space'), '360')  # ~0.25"
cols.set(qn('w:equalWidth'), '1')

# ----- Helpers -----
def add_body(text, justify=True, first_indent=0.2):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(first_indent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    return p

def add_lead_para(prefix_bold_italic, body_italic_text):
    """Used for Abstract and Index Terms lines."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(prefix_bold_italic)
    r1.bold = True
    r1.italic = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(9)
    r2 = p.add_run(body_italic_text)
    r2.italic = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(9)
    return p

def add_primary_heading(text):
    """Centered, small caps, 10pt, with extra space."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    rPr = r._element.get_or_add_rPr()
    sc = OxmlElement('w:smallCaps')
    sc.set(qn('w:val'), '1')
    rPr.append(sc)
    return p

def add_secondary_heading(text):
    """Italic, left-aligned letter heading (A., B., ...)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    return p

def add_reference(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(8)
    return p

# ----- Abstract & Index Terms -----
add_lead_para("Abstract—", abstract_text)
add_lead_para("Index Terms—", keywords_text)

# ----- I. Introduction -----
add_primary_heading("I.  Introduction")
for para in intro_body:
    add_body(para)

# ----- II. Deployment Context (groups the 2.x.x sections) -----
add_primary_heading("II.  Deployment Context and Design Environment")

add_secondary_heading("A.  Deployment Environments and Users")
for para in sub_a_body:
    add_body(para)

add_secondary_heading("B.  Environmental Density and Privacy Considerations")
for para in sub_b_body:
    add_body(para)

add_secondary_heading("C.  Environmental Challenges and Design Responses")
for para in sub_c_body:
    add_body(para)

# ----- III. Ethical Implications -----
add_primary_heading("III.  Ethical Implications of Commercialization")
for para in eth_body:
    add_body(para)

# ----- IV. Lessons Learned -----
add_primary_heading("IV.  Professional and Ethical Responsibility: Lessons Learned")
for para in les_body:
    add_body(para)

# ----- V. Impact -----
add_primary_heading("V.  Global, Economic, Environmental, and Societal Impact")
for para in imp_body:
    add_body(para)

# ----- VI. Conclusion -----
add_primary_heading("VI.  Conclusion")
for para in con_body:
    add_body(para)

# ----- References (unnumbered primary heading per IEEE) -----
add_primary_heading("References")
for ref in references:
    add_reference(ref)

# NOTE: do NOT strip em-dashes here — the "Abstract—" and "Index Terms—"
# prefixes are required IEEE formatting. Body text was already cleaned earlier.

doc.save(SRC)
print('IEEE-formatted essay saved.')
