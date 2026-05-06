import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

doc = docx.Document('essay1_BMCS.docx')
text_blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# Remove any empty blocks
text_blocks = [t for t in text_blocks if t]

# The first 4 lines in the original are the title and author
# text_blocks[0]: "Understanding Professional and Ethical Responsibility"
# text_blocks[1]: "and Global, Economic, Environmental and Societal Impact"
# text_blocks[2]: "of the BMCS Project"
# text_blocks[3]: "Kaushik Vada   |   EE175AB Senior Design   |   UC Riverside   |   May 5, 2026"
# text_blocks[4]: "Ethical Implications of Commercialization"
# text_blocks[5]: "Battery technology sits at the center..." (Introduction)

# Find indices
try:
    ethical_idx = text_blocks.index("Ethical Implications of Commercialization")
except ValueError:
    ethical_idx = 4

try:
    global_idx = text_blocks.index("Global, Economic, and Environmental Impact")
except ValueError:
    global_idx = -1

title = "Understanding Professional and Ethical Responsibility and Global, Economic, Environmental and Societal Impact of the BMCS Project"
author_info = "Kaushik Vada   |   EE175AB Senior Design   |   UC Riverside   |   May 5, 2026"

intro_text = [text_blocks[ethical_idx + 1]]
ethical_text = text_blocks[ethical_idx + 2 : global_idx]
global_text = text_blocks[global_idx + 1 : -1] # Exclude the last paragraph which is Conclusion
conclusion_text = [text_blocks[-1]]

new_doc = docx.Document()

# Title
title_p = new_doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run(title)
title_run.bold = True
title_run.font.size = Pt(14)

# Author
author_p = new_doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_p.add_run(author_info).font.size = Pt(11)

# Abstract
abstract_heading = new_doc.add_paragraph()
abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract_heading.add_run("Abstract").bold = True

abstract_text = (
    "This essay explores the professional and ethical responsibilities, alongside the global, "
    "economic, and environmental impacts, inherent in the design and potential commercialization "
    "of a Battery Management and Control System (BMCS). Built around a 10s1p Molicel P42A "
    "lithium-ion pack, the prototype highlights critical ethical obligations prioritizing public "
    "safety over cost and complexity. By analyzing the catastrophic consequences of inadequate "
    "cell monitoring—demonstrated by historical failures like the Samsung Galaxy Note 7 and "
    "Boeing 787—the essay emphasizes the severe asymmetry in knowledge between engineers and consumers. "
    "The analysis further delves into the privacy concerns of telemetry data logging and the "
    "transparency required in commercial deployments. From a global perspective, the essay discusses "
    "how accessible, reliable battery management is a fundamental enabler of the broader energy "
    "transition, supporting applications from e-bikes to off-grid solar installations in developing "
    "regions. The environmental duality of lithium-ion technology is also examined, acknowledging "
    "both the benefits of extended cell cycle life and the ecological costs of resource extraction. "
    "Ultimately, this work concludes that strict adherence to standards such as UL 9540 and "
    "IEC 62619 is not merely bureaucratic but a fundamental engineering responsibility."
)
abs_p = new_doc.add_paragraph(abstract_text)
abs_p.paragraph_format.left_indent = Inches(0.5)
abs_p.paragraph_format.right_indent = Inches(0.5)
for run in abs_p.runs:
    run.font.italic = True

# Keywords
keywords_p = new_doc.add_paragraph()
keywords_p.paragraph_format.left_indent = Inches(0.5)
keywords_p.paragraph_format.right_indent = Inches(0.5)
kw_run = keywords_p.add_run("Index Terms—")
kw_run.bold = True
kw_run.font.italic = True
keywords_p.add_run("Battery Management Systems (BMS), Engineering Ethics, Public Safety, Renewable Energy Storage, Lithium-ion Batteries.")

def add_heading(text, level=1):
    h = new_doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    if level == 1:
        r.font.size = Pt(10)

add_heading("I. INTRODUCTION")
for t in intro_text:
    new_doc.add_paragraph(t)

add_heading("II. ETHICAL IMPLICATIONS OF COMMERCIALIZATION")
for i, t in enumerate(ethical_text):
    if "Samsung Galaxy Note 7" in t and "Boeing 787" in t:
        t = t.replace("Boeing 787 Dreamliner groundings", "Boeing 787 Dreamliner groundings [1], [2]")
    new_doc.add_paragraph(t)

add_heading("III. GLOBAL, ECONOMIC, AND ENVIRONMENTAL IMPACT")
for t in global_text:
    if "International Energy Agency has noted" in t:
        t = t.replace("International Energy Agency has noted", "International Energy Agency (IEA) has noted [3]")
    if "UL 9540 and IEC 62619" in t:
        t = t.replace("UL 9540 and IEC 62619", "UL 9540 [4] and IEC 62619 [5]")
    new_doc.add_paragraph(t)

add_heading("IV. CONCLUSION")
for t in conclusion_text:
    new_doc.add_paragraph(t)

add_heading("V. REFERENCES")
refs = [
    '[1] Consumer Product Safety Commission, "Samsung Galaxy Note 7 Battery Fire Investigation," CPSC, Washington, DC, Rep. 2017.',
    '[2] National Transportation Safety Board (NTSB), "Auxiliary Power Unit Battery Fire Japan Airlines Boeing 787-8, JA829J," NTSB, Washington, DC, Rep. AIR-14-01, 2014.',
    '[3] International Energy Agency (IEA), "SDG7: Data and Projections - Access to Electricity," IEA, Paris, 2022. [Online]. Available: https://www.iea.org/reports/sdg7-data-and-projections/access-to-electricity.',
    '[4] UL Standard for Safety for Energy Storage Systems and Equipment, UL 9540, 2023.',
    '[5] Secondary cells and batteries containing alkaline or other non-acid electrolytes - Safety requirements for secondary lithium cells and batteries, for use in industrial applications, IEC 62619:2022, 2022.',
    '[6] K. Vada et al., "Battery Management System," EE175AB Final Report, Dept. Elect. and Comput. Eng., Univ. California, Riverside, CA, USA, 2026.'
]

for ref in refs:
    p = new_doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.add_run(ref).font.size = Pt(9)

new_doc.save("version 2 - Essay 1.docx")
print("Successfully generated v2")
