"""
Build Ethics_Impact_Essay_v2.docx in IEEE journal format.
Title/abstract span full width; body uses a borderless two-column table.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell

FONT    = 'Times New Roman'
COL_W   = 3.525   # each column in inches  (7.25 total - 0.2 gap) / 2
GAP_TWP = 144     # gap between cols in twips (0.1 in each side = 72 twips)


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _run(p, text, size=10, bold=False, italic=False, sc=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold   = bold
    r.italic = italic
    if sc:
        rPr = r._r.get_or_add_rPr()
        rPr.append(OxmlElement('w:smallCaps'))
    return r


def _fmt(p, align, sb, sa, li, ri, fi):
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after  = Pt(sa)
    if li is not None: pf.left_indent       = Inches(li)
    if ri is not None: pf.right_indent      = Inches(ri)
    if fi is not None: pf.first_line_indent = Inches(fi)


# ---------------------------------------------------------------------------
# Writer  —  wraps Document or _Cell, handles the auto-empty first paragraph
# ---------------------------------------------------------------------------

class W:
    def __init__(self, container):
        self._c     = container
        self._first = isinstance(container, _Cell)   # cells have a starter para

    def _next(self):
        if self._first:
            self._first = False
            return self._c.paragraphs[0]
        return self._c.add_paragraph()

    def para(self, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             sb=0, sa=4, li=None, ri=None, fi=None):
        p = self._next()
        _fmt(p, align, sb, sa, li, ri, fi)
        return p

    def body(self, text):
        p = self.para(sb=0, sa=3, fi=0.18)
        _run(p, text)

    def heading(self, text):
        p = self.para(align=WD_ALIGN_PARAGRAPH.CENTER, sb=8, sa=4)
        _run(p, text, sc=True)

    def ref(self, num, text):
        p = self.para(sb=0, sa=2, li=0.22, fi=-0.22)
        _run(p, num + ' ', size=8)
        _run(p, text, size=8)


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _no_borders(table):
    """Remove every visible border from a table and its cells."""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tb = OxmlElement('w:tblBorders')
    for s in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{s}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tb.append(el)
    tblPr.append(tb)

    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            cb   = OxmlElement('w:tcBorders')
            for s in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                el = OxmlElement(f'w:{s}')
                el.set(qn('w:val'),   'none')
                el.set(qn('w:sz'),    '0')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), 'auto')
                cb.append(el)
            tcPr.append(cb)


def _set_table_layout(table, col_widths_in):
    """Set fixed table width and column grid."""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))

    # Fixed total width
    tblW = OxmlElement('w:tblW')
    total = sum(col_widths_in) * 1440
    tblW.set(qn('w:w'),    str(int(total)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # tblGrid column definitions
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths_in:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w * 1440)))
        tblGrid.append(gc)
    tbl.insert(tbl.index(tblPr) + 1, tblGrid)


def _cell_width(cell, inches):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _cell_margins(cell, left=0, right=0, top=0, bot=0):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    m = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bot),
                      ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.top_margin    = Inches(0.75)
sec.bottom_margin = Inches(1.0)
sec.left_margin   = Inches(0.625)
sec.right_margin  = Inches(0.625)

# Ensure single-column main section (no section-break tricks)
sp = sec._sectPr
for c in sp.findall(qn('w:cols')):
    sp.remove(c)

d = W(doc)

# ── Title ──────────────────────────────────────────────────────────────────
p = d.para(align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
_run(p, 'Understanding of Professional and Ethical Responsibility and Global,\n'
        'Economic, Environmental, and Societal Impact', size=24)

p = d.para(align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=6)
_run(p, 'Modular Battery Management and Communication Platform', size=16)

p = d.para(align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=2)
_run(p, 'Kaushik Vada, Connor Stewart, Nick Poon, Christian Kim', size=11)

p = d.para(align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=8)
_run(p, 'UC Riverside, EE 175AB, 2026', size=11, italic=True)

# ── Two-column body (borderless table) ────────────────────────────────────
table = doc.add_table(rows=1, cols=2, style='Table Grid')
_no_borders(table)
_set_table_layout(table, [COL_W, COL_W])

L, R = table.cell(0, 0), table.cell(0, 1)
_cell_width(L, COL_W);  _cell_margins(L, left=0,       right=GAP_TWP)
_cell_width(R, COL_W);  _cell_margins(R, left=GAP_TWP, right=0)

lc = W(L)
rc = W(R)

# ══════════════════════ LEFT COLUMN ═══════════════════════════════════════


# Abstract and Index Terms inside the left column (matches IEEE layout)
p = lc.para(sa=5)
_run(p, 'Abstract—', size=9, bold=True, italic=True)
_run(p,
    'This essay examines the professional, ethical, environmental, and societal '
    'implications of the Modular Battery Management and Communication Platform, a '
    '36 V lithium-ion BMS developed for the EE 175AB senior design sequence at '
    'UC Riverside. The system supplies up to 4 A from a 10-series Molicel P42A '
    'pack, targeting personal electric mobility and embedded robotics. A commercial '
    'version would operate in dense urban environments where lithium-ion fires have '
    'already caused fatalities and prompted new certification rules in New York City '
    'and the European Union. The essay covers public safety and privacy, how the '
    'current design addresses them through layered hardware and firmware safeguards, '
    'what the team learned about engineering honesty, and the likely effects of '
    'low-cost open-architecture battery management on global markets.',
    size=9, italic=True)

p = lc.para(sa=10)
_run(p, 'Index Terms—', size=9, bold=True, italic=True)
_run(p,
    'battery management system, lithium-ion safety, micro-mobility, '
    'telemetry privacy, IEEE Code of Ethics',
    size=9, italic=True)

lc.heading('I.  Introduction')
lc.body(
    'Lithium-ion fires in electric scooters and e-bikes have caused deaths, '
    'structural damage, and legislation in New York City, the European Union, '
    'and elsewhere [2]. At the same time, electrifying two-wheeled transport is '
    'one of the cheaper ways to cut transport-sector emissions in growing cities '
    '[7]. Engineers designing battery management systems sit between those two '
    'pressures, and there is no clean way to resolve them except to take safety '
    'seriously from the first schematic.')
lc.body(
    'The Modular Battery Management and Communication Platform manages a '
    '10-series lithium-ion pack with hardware-enforced fault protection and '
    'real-time telemetry. Its modular architecture lets the same board adapt to '
    'electric scooters, e-bikes, and embedded robotics platforms. The prototype '
    'was tested in a university lab, but a commercial version would run in '
    'public spaces where a single design failure can hurt bystanders who had '
    'nothing to do with the product.')
lc.body(
    'This essay addresses four questions: what ethical issues a commercial '
    'version raises, how the current design responds to them, what the team '
    'learned about professional responsibility while building it, and what '
    'broader deployment would likely mean for the world.')

lc.heading('II.  Deployment Context')
lc.body(
    'The BMS outputs 36 V nominal at up to 4 A from a 10-series Molicel '
    'P42A pack. The prototype was tested in a controlled lab at UC Riverside, but a commercial version would operate in far less controlled environments.')
lc.body(
    'Shared scooter and e-bike fleets operate in cities like Los Angeles, '
    'Amsterdam, Shenzhen, and Nairobi, each with different climate profiles, '
    'including tropical humidity, monsoon rain, and sub-freezing winters [1]. '
    'The BMS would also power embedded platforms like the NVIDIA Jetson Orin '
    'Nano in field robotics settings, where motor-drive EMI and sustained '
    'vibration are routine. End users range from riders who have never opened '
    'a battery enclosure to research engineers who read the telemetry GUI '
    'without knowing much about what the hardware underneath is doing. That '
    'range is what the design team is actually responsible for.')
lc.body(
    'The pack stores about 151 Wh. When something goes wrong with a pack '
    'that size on a crowded sidewalk or in a transit station, the consequences '
    'land on people who did not choose to be near it. That context drove every '
    'safety decision in the design.')
lc.body(
    'On data privacy: the BMS currently sends telemetry over USB CDC in a '
    'closed lab environment. In a commercial fleet product, the same data sent '
    'wirelessly carries location and usage patterns that fall under CCPA [3] '
    'and GDPR [4]. Engineers designing that version need to plan for '
    'anonymization, encryption, and access controls before the data starts '
    'flowing, not after.')

lc.heading('III.  Ethical Implications of Commercialization')
lc.body(
    'Lithium battery fires caused 18 deaths in New York City in 2023, '
    'which led to local legislation requiring UL '
    'certification for lithium batteries sold or rented in the city [2]. A '
    'BMS that goes to market has to protect reliably across the full service '
    'life of the pack, not just at nominal conditions on a bench.')
lc.body(
    'The design uses hardware-layered protection for this reason. The Texas '
    'Instruments BQ76930 analog front end [12] handles fault protection '
    'independently of the STM32 firmware. Short-circuit protection trips '
    'within 70 µs with no software involvement, consistent with IEC 62133-2 short-circuit and overcharge requirements for portable lithium cells [5], since firmware-dependent '
    'protection is exposed to software bugs and communication failures that '
    'have contributed to real incidents. The mechanical relay serving as the '
    'final breaker replaced the original MOSFET-based design after datasheet '
    'discrepancies showed it could not be relied on under certain fault '
    'conditions. The change was made under schedule and budget constraints, '
    'consistent with the IEEE Code of Ethics requirement that engineers '
    'hold paramount the safety of the public [6].')
lc.body(
    'Open architecture is also an ethical choice in this context. The '
    'first-run prototype cost about approximately $3,000 including development hardware. '
    'Separating the battery stack, firmware, and visualization layer makes '
    'the design easier to cost-reduce at volume and easier to adapt for '
    'different cell configurations. Hosting schematics and firmware on GitHub '
    'makes that knowledge available to student teams and small operators who '
    'cannot afford commercial alternatives.')
lc.body(
    'The telemetry pipeline carries a data-ethics obligation even in its '
    'current form. Data minimization, purpose limitation, and user consent '
    'are not features that can be added later without restructuring what '
    'already exists. If the team ever builds a wireless fleet version, those '
    'decisions should already be made.')

# ══════════════════════ RIGHT COLUMN ══════════════════════════════════════

rc.heading('IV.  Professional and Ethical Responsibility: Lessons Learned')
rc.body(
    'The principal design lesson was that fault handling cannot be retrofitted; it must be designed in from the schematic stage. The IEEE Code of Ethics '
    'says to hold public safety paramount [6]. In practice, that meant every '
    'fault condition, including overcurrent, overvoltage, undervoltage, short '
    'circuit, and overtemperature, required documented test cases at 100% '
    'pass rate before the BMS connected to a live pack. The Cell Emulator '
    'PCB was built specifically so hardware fault testing did not expose the '
    'team to a live lithium pack during validation. Protecting the people '
    'who test the system matters as much as protecting the people who use it.')
rc.body(
    'A second lesson concerned honest reporting when measured '
    'performance fell short of design targets. Post-production routing errors in the E-Load PCB made '
    'the original 5 A target unachievable. The options were to report '
    'numbers that could not be backed up, or to document the limitation, '
    'derate to 1 A, and describe the fix. The IEEE Code of Ethics asks '
    'engineers to be honest in stating claims based on available data [6]. '
    'In a safety-critical system, overstating performance is not just an '
    'embarrassment, it is a hazard, because users design their own systems '
    'around those specifications.')
rc.body(
    'Documentation is engineering work, not overhead. Version-controlled '
    'PCB files, a firmware repository with a real README, and operator '
    'installation guides are how knowledge about a safety-critical system '
    'survives the engineers who built it. That matters more in this domain '
    'than in most.')

rc.heading('V.  Global, Economic, Environmental, and Societal Impact')
rc.body(
    'The biggest global opportunity for this platform is in developing '
    'markets. Across sub-Saharan Africa, South and Southeast Asia, and Latin '
    'America, urban mobility runs on petrol-powered motorcycles and scooters '
    'that are a major source of local air pollution and transport-sector '
    'carbon. The IEA treats electrifying those fleets as one of the more '
    'accessible mitigation paths [7]. The constraint is not the cells or the '
    'motors, it is battery management that local manufacturers can actually '
    'source, repair, and understand. An open-architecture BMS built to '
    'IPC-2221 standards [10] with AEC-Q components [11] where possible is '
    'useful in that context in a way that a closed commercial unit is not.')
rc.body(
    'Cell balancing also matters environmentally at scale. Keeping weaker '
    'cells from degrading faster than the rest of the pack extends total '
    'pack life, which means fewer packs manufactured, less lithium and '
    'cobalt extracted, and less battery waste. That benefit compounds with '
    'adoption.')
rc.body(
    'Regulatorily, the EU Battery Regulation (EU) 2023/1542 [8] now '
    'requires BMS performance data, carbon footprint declarations, and '
    'end-of-life collection for EV batteries including light electric '
    'vehicles. The CPSC is developing mandatory safety standards for '
    'e-mobility devices in the US after a series of high-profile fires [9]. '
    'Early conformance to IPC-2221, hardware fault protection, and clean '
    'signal integrity should make certification tractable when the time '
    'comes.')
rc.body(
    'In rural and peri-urban communities with weak grid access, reliable '
    'battery storage supports small business, medical refrigeration, and '
    'communications infrastructure. A BMS with separated, documented layers '
    'is easier to adapt and teach where engineering budgets are small and '
    'staff learns by doing.')

rc.heading('VI.  Conclusion')
rc.body(
    'Technically, this is a 36 V lithium-ion safety and telemetry system. '
    'The harder part was the judgment calls: what to protect, how to handle '
    'a spec that could not be met, what to put in writing. '
    'Hardware-independent fault protection, honest performance reporting, '
    'open-source architecture, conformance to industry standards, and '
    'documented fault validation are all consistent with what the IEEE Code '
    'of Ethics actually asks for. The global and societal potential of '
    'affordable, safe battery management is real. Whether it matters in '
    'practice depends on whether the engineers who build and sell it treat '
    'public safety as the first constraint in the design rather than a '
    'late-stage compliance check.')

rc.heading('References')
REFS = [
    ('[1]',  'Bird Global, Inc., “Sustainability Report,” Miami, FL, USA, 2022. [Online]. Available: https://www.bird.co/sustainability/'),
    ('[2]',  'New York City Fire Department, “FDNY Public Safety Bulletin: Lithium-Ion Battery Fires,” New York, NY, USA, 2023. [Online]. Available: https://www.nyc.gov/site/fdny/'),
    ('[3]',  'State of California, “California Consumer Privacy Act of 2018,” Cal. Civ. Code, sec. 1798.100 et seq., 2018.'),
    ('[4]',  'European Parliament and Council, “Regulation (EU) 2016/679 (General Data Protection Regulation),” Official Journal of the European Union, vol. L 119, pp. 1–88, May 2016.'),
    ('[5]',  'International Electrotechnical Commission, “IEC 62133-2: Secondary cells and batteries containing alkaline or other non-acid electrolytes. Safety requirements for portable sealed secondary lithium cells,” IEC, Geneva, Switzerland, 2017.'),
    ('[6]',  'IEEE, “IEEE Code of Ethics,” IEEE Policies, sec. 7.8, 2020. [Online]. Available: https://www.ieee.org/about/corporate/governance/p7-8.html'),
    ('[7]',  'International Energy Agency, “Global EV Outlook 2024,” IEA, Paris, France, 2024. [Online]. Available: https://www.iea.org/reports/global-ev-outlook-2024'),
    ('[8]',  'European Parliament and Council, “Regulation (EU) 2023/1542 concerning batteries and waste batteries,” Official Journal of the European Union, vol. L 191, pp. 1–117, Jul. 2023.'),
    ('[9]',  'U.S. Consumer Product Safety Commission, “Letter to Manufacturers, Importers, Distributors, and Retailers of Lithium-Ion Batteries and Products Containing Such Batteries,” CPSC, Bethesda, MD, USA, Dec. 2022.'),
    ('[10]', 'IPC, “IPC-2221B: Generic Standard on Printed Board Design,” IPC, Bannockburn, IL, USA, 2012.'),
    ('[11]', 'Automotive Electronics Council, “AEC-Q100 Rev-H: Failure Mechanism Based Stress Test Qualification for Integrated Circuits,” AEC, 2014.'),
    ('[12]', 'Texas Instruments, “BQ76930 3-Series to 15-Series Cell Battery Monitor and Protector,” Datasheet SLUSBK3, Dallas, TX, USA, 2018. [Online]. Available: https://www.ti.com/product/BQ76930'),
]
for num, text in REFS:
    rc.ref(num, text)

doc.save('Ethics_Impact_Essay_v3.docx')
print('Saved Ethics_Impact_Essay_v3.docx')
