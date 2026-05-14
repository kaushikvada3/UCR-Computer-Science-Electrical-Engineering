"""Apply humanized text to Ethics_Impact_Essay.docx, preserving paragraph/run formatting."""
from docx import Document
from copy import deepcopy

PATH = r"C:\Users\chris_bptnrfq\Desktop\181W\ENGR181W - Technical Writing Class\Ethics_Impact_Essay.docx"

# Map of paragraph index -> new text. Indices reference the original document.
# IEEE-style abstract/index-terms keep their "Abstract—" / "Index Terms—" prefix in the first run for italic formatting.

NEW = {}

NEW[5] = (
    "Abstract—",
    "This essay examines the professional, ethical, environmental, and societal implications of the Modular Battery Management and Communication Platform, a 36 V lithium-ion battery management system (BMS) developed for the EE 175AB senior design sequence at UC Riverside. The platform supplies up to 4 A from a 10-series Molicel P42A pack and is built for personal electric mobility and embedded robotics. A commercial version would operate in dense urban settings where lithium-ion fires have already caused fatalities and prompted new certification rules in New York City and the European Union Battery Regulation 2023/1542. The paper covers four areas: the public safety, privacy, and equity implications of deployment; how the present design handles each through layered hardware and firmware safeguards; what the team learned about engineering honesty and documentation; and the likely global, environmental, and societal effects of low-cost open-architecture battery management. Responsible commercialization depends on treating public safety as a first commitment in the design, not a late constraint."
)

NEW[6] = (
    "Index Terms—",
    "battery management system, lithium-ion safety, micro-mobility, telemetry privacy, IEEE Code of Ethics"
)

NEW[8] = (
    None,
    "Lithium-ion energy storage has grown quickly worldwide, and so have the safety incidents, regulatory actions, and public concerns that follow it. Fires in electric scooters, e-bikes, and similar devices have caused injuries, deaths, and structural damage in major cities, and have led to new certification rules and import restrictions [2]. At the same time, electrifying transport remains one of the cheaper ways to cut greenhouse gas emissions in fast-growing cities [7]. Engineers who design battery management systems work between two pressures at once: the push to deploy energy storage broadly, and the duty to ensure each unit is safe for the people around it."
)

NEW[9] = (
    None,
    "The Modular Battery Management and Communication Platform is a senior design project aimed at this gap. The system manages a 10-series lithium-ion pack with hardware-enforced fault protection and real-time telemetry, and uses a modular architecture so the same design can be adapted to electric scooters, e-bikes, or embedded robotics platforms. The prototype was built in a controlled lab, but a commercial version would run in places where one design failure can hurt people who are not the rider."
)

NEW[10] = (
    None,
    "The essay answers four questions: (a) what ethical issues a commercial version of this system would raise; (b) how the present design responds to them; (c) what the team learned about professional and ethical responsibility while building it; and (d) what global, economic, environmental, and societal effects broader deployment would likely have. The sections that follow cover deployment context, ethical implications, lessons learned, broader impact, and a conclusion."
)

NEW[13] = (
    None,
    "The BMS is a compact, plug-and-play power source for mid-sized electronics. It supplies a nominal 36 V at up to 4 A from a 10-series Molicel P42A lithium-ion cell pack. The prototype was built and tested in a university laboratory at UC Riverside, but a commercial version would have to work in a much wider and harder set of physical conditions."
)

NEW[14] = (
    None,
    "The main application is personal electric mobility, specifically electric scooters and e-bikes. That puts the BMS in outdoor, urban, and suburban settings worldwide. In the United States, operators like Bird and Lime have deployed hundreds of thousands of shared scooters across dozens of cities, including Los Angeles, New York, and smaller college towns, where the packs see sidewalks, bike lanes, transit hubs, and parking structures [1]. Outside the United States, e-scooters and e-bikes have become primary transport in dense cities across Europe, Southeast Asia, and South America. Two-wheeled electric mobility has grown rapidly in cities such as Amsterdam, Shenzhen, Bogotá, and Nairobi, so a commercial version of this BMS could see tropical humidity, monsoon rain, sub-freezing winters, and high-altitude operation over a single product lifetime."
)

NEW[15] = (
    None,
    "A second use case the team identified is powering embedded computing platforms, such as the NVIDIA Jetson Orin Nano, for robotics and autonomous systems research. The BMS would then sit in laboratory and field robotics settings: research institutions, manufacturing floors, and outdoor autonomous vehicle ranges. These environments add vibration, electromagnetic interference from motor drives, and duty cycles that differ from consumer scooter use."
)

NEW[16] = (
    None,
    "End users vary widely. Consumer products are used by everyday riders, often young adults and students with no formal electronics training. In shared fleets, the battery system is handled by technicians who need clear fault indicators and a straightforward maintenance interface. In robotics and research, the users are engineers and graduate students who know their domain but often rely on the telemetry GUI for real-time diagnostics without deep knowledge of the hardware underneath. That range of users shapes what the design team is responsible for."
)

NEW[18] = (
    None,
    "Electric scooters are everywhere in crowded urban areas: sidewalks, transit stations, college campuses, downtown business districts. The relevance to the BMS is mostly about public safety. The pack stores about 151 Wh of energy, so a thermal runaway does not happen in isolation; it happens in a crowd. An uncontrolled fire in a public space can injure bystanders, damage property, and erode public trust in electric mobility."
)

NEW[19] = (
    None,
    "On data privacy, the BMS as designed sends a 100-byte telemetry frame (cell voltages, temperatures, current draw, and fault status) over USB CDC to a connected host computer. In a laboratory the data path is closed and benign. In a commercial fleet management product, the same telemetry could be sent wirelessly to a cloud backend for fleet health monitoring, and at that point it becomes a privacy concern. GPS-tagged battery telemetry can reveal a rider's location history, usage patterns, and behavior. Engineers building such systems are responsible for anonymizing the data, encrypting it in transit, and storing it under access controls that meet rules like the California Consumer Privacy Act (CCPA) [3] and the European General Data Protection Regulation (GDPR) [4]."
)

NEW[21] = (
    None,
    "The physical environment creates engineering problems that touch directly on safety. Lithium-ion cells are sensitive to temperature extremes. High ambient temperatures speed up capacity fade and raise the risk of thermal runaway; low temperatures cut available capacity and can pull cells below their safe minimum voltage if discharge continues. The BMS handles this with continuous NTC thermistor monitoring on every thermal channel, a PWM-controlled fan for active cooling, and a hardware-enforced breaker that disconnects the load above 60 °C. The team treated these as primary safety requirements rather than optional features."
)

NEW[22] = (
    None,
    "Moisture and ingress protection matter for any outdoor deployment. The current prototype runs in a laboratory enclosure and has no IP (Ingress Protection) rating. A commercial scooter product would need at least IP54 for splash resistance, and preferably IP65 for sustained outdoor use. In the European Union, the Low Voltage Directive (LVD) and EN 62133 for secondary lithium cells [5] set specific thermal, short-circuit, and overcharge requirements that match the fault protections already in this design."
)

NEW[23] = (
    None,
    "Electromagnetic interference (EMI) matters in motor-drive settings. The team routed the USB communication lines as differential pairs to meet USB 2.0 Full-Speed signal integrity requirements and added ESD protection diodes on every external-facing port because electrical noise is the norm in the environments this product is built for. Those choices were made so the system stays reliable and safe under realistic operating conditions."
)

NEW[25] = (
    None,
    "Public safety is the most important ethical issue here. Lithium-ion battery fires in consumer products such as e-scooters, e-bikes, and hoverboards have killed and injured people and started fires in residential buildings. In New York City alone, lithium battery fires caused more than a dozen deaths in a single year and led to emergency legislation requiring UL certification for lithium batteries sold or rented in the city [2]. A BMS sold to the public has to perform its protective functions reliably in real conditions, for the full service life of the pack."
)

NEW[26] = (
    None,
    "The design uses a layered safety architecture for this reason. At the hardware level, the Texas Instruments BQ76930 analog front end [12] handles fault protection independently of the STM32 firmware: short-circuit protection trips within 70 ms with no software involvement. This matters because firmware-dependent protection is exposed to software bugs, stack overflows, and communication failures, and all three have contributed to real-world battery incidents. Adding a redundant mechanical relay as the final breaker, picked to meet the current and voltage ratings of the 10s lithium pack, was a deliberate safety choice rather than an afterthought. The team replaced the original MOSFET-based breaker after datasheet discrepancies risked leaving the system exposed, and put in a validated mechanical relay to close the gap. The decision was made under time and budget pressure, but it is the kind of judgment the IEEE Code of Ethics asks for: hold public safety paramount even when it is inconvenient."
)

NEW[27] = (
    None,
    "Fairness of access is a smaller but real concern. The BMS was built to replace expensive commercial units with a low-cost, open-architecture alternative. The first-run prototype cost about $2,966 including development hardware, and the modular architecture separates the battery stack, the firmware, and the visualization layer, which makes the cost easier to bring down at volume and the system easier to adapt to different cell configurations. Reliable battery management that small-scale EV builders, student robotics teams, and low-income micro-mobility operators in developing markets can actually afford is a real societal benefit. Hosting the firmware and GUI openly on GitHub follows the same engineering principle of sharing knowledge with the profession."
)

NEW[28] = (
    None,
    "The real-time telemetry is mainly a safety and diagnostic tool, but it carries data-ethics obligations once the system is sold commercially. The PyQt6 GUI currently shows telemetry locally over USB, a closed path with no privacy risk. The same pipeline could be extended to a wireless, cloud-connected version for a fleet management product. At that point the telemetry log becomes a record of user behavior, and the designers are responsible for data minimization, purpose limitation, and user consent under the applicable privacy law. Engineers who build systems that produce sensitive data should plan for these second-order uses before the product ships."
)

NEW[30] = (
    None,
    "The project taught practical lessons in professional and ethical responsibility that the classroom could not. The biggest one was designing for failure modes. The IEEE Code of Ethics tells engineers to “hold paramount the safety, health, and welfare of the public” [6]. In practice that meant the team could not just verify the BMS under nominal conditions and call it done. Each fault condition (overcurrent, overvoltage, undervoltage, short circuit, and overtemperature) had to be validated with documented test cases at a 100% pass rate before the BMS could be connected to a live lithium pack. The Cell Emulator PCB was built so that hardware fault testing did not put the team or the lab at risk from a live pack, and it is itself an ethical artifact: responsible engineering practice protects the people who validate the system, not only the people who use it."
)

NEW[31] = (
    None,
    "A second lesson was about revising a design under resource constraints. Post-production routing errors and component discrepancies in the E-Load PCB made the original 5 A target impossible. The team had two options: report numbers we could not back up, or document the limitation, derate the spec to 1 A, and lay out the path to fix it later. We chose the second. The IEEE Code of Ethics asks engineers to “be honest and realistic in stating claims or estimates based on available data” [6]. In industry, overstating the performance of a safety-critical system is more than embarrassing; users rely on those specifications, and the gap between claim and reality can contribute to accidents."
)

NEW[32] = (
    None,
    "The team also learned that documentation is part of professional responsibility. Keeping the PCB files in version control, a firmware repository with a real README, and release documentation that includes operator installation guides is not paperwork. It is how engineering knowledge gets preserved and reviewed, and it is what the next engineer needs to maintain or modify the system safely. Poor documentation of safety-critical systems has contributed to engineering disasters in fields ranging from industrial process control to medical devices. Documenting properly is part of the engineering work, not extra to it."
)

NEW[34] = (
    None,
    "If commercialized, the largest global impact of this platform would be in helping safe, affordable electric micro-mobility take hold in the developing world. In cities across sub-Saharan Africa, South and Southeast Asia, and Latin America, micro-mobility is dominated by petrol-powered motorcycles and scooters that add a lot of urban air pollution and greenhouse gas emissions. The International Energy Agency [7] treats electrifying these fleets as one of the cheaper ways to cut transport-sector carbon in fast-growing cities. A key barrier is the cost and availability of reliable battery management technology, because unsafe or unreliable packs are both a safety hazard and a reason not to switch. A low-cost, open-architecture BMS built to automotive-grade safety standards (the team followed IPC-2221 clearance and creepage requirements [10] and used AEC-Q qualified components where possible [11]) could lower that barrier and let local entrepreneurs and small-scale manufacturers build safer electric vehicles for their own markets."
)

NEW[35] = (
    None,
    "Environmentally, the project touches the lithium-ion battery lifecycle, which is under growing regulatory and ethical scrutiny. Lithium-ion cells contain lithium, cobalt, and nickel, whose extraction carries serious environmental and human-rights costs where they are mined. Extending battery service life through precise cell management therefore has both an economic and an environmental case: every extra charge cycle a pack delivers under accurate cell balancing and fault-protected operation pushes back the environmental cost of building a replacement. The cell balancing in this BMS equalizes cell voltages to keep the weakest cells in a series stack from losing capacity early, which extends pack life and produces a measurable environmental benefit at commercial scale."
)

NEW[36] = (
    None,
    "There is also active regulatory and political activity in this area. In the United States, the Consumer Product Safety Commission (CPSC) has issued guidance on lithium battery safety [9] and is developing mandatory standards for e-mobility devices after a string of high-profile fires. In the European Union, the Battery Regulation (EU) 2023/1542, which came into force in August 2023 [8], sets mandatory requirements for battery management system performance, carbon footprint declarations, and end-of-life collection, and these apply to EV batteries including those in light electric vehicles. A commercial version of this BMS would have to be designed for those rules. The team's early use of IPC-2221, USB 2.0 Full-Speed, and hardware-level fault protection should make later certification easier."
)

NEW[37] = (
    None,
    "The societal effect of reliable battery management also reaches energy equity. In rural and peri-urban communities in the developing world, weak grid access makes battery storage essential for everyday economic activity: small businesses, medical refrigeration, and communication equipment. A BMS that can be adapted to different cell configurations, capacities, and host microcontrollers, and whose firmware and schematics are openly available, helps spread energy storage technology beyond the companies that already have it. The modular architecture here, with the battery stack, firmware, and visualization layers separated, was chosen for engineering reasons, but it also makes the technology easier to adapt and teach in places without large engineering budgets."
)

NEW[39] = (
    None,
    "Technically, the Modular Battery Management and Communication Platform is a 36 V lithium-ion safety and telemetry system. Ethically, it is an exercise in professional responsibility applied to rechargeable energy storage, where a design failure is measured in fires, injuries, and lost public trust. The design choices in this project (hardware-independent fault protection, honest performance documentation, open-source architecture, conformance to industry standards, and thorough testing against failure modes) match the obligations the IEEE Code of Ethics places on practicing engineers. The global, environmental, and societal potential of this technology is real, but it only shows up when the engineers who build and deploy it treat public safety as the first commitment, not as a constraint to push down."
)


def replace_paragraph_text(p, prefix_text, body_text):
    """Replace paragraph text. If prefix_text given, keep first run formatting for prefix
    (e.g. italic 'Abstract—') and put body in a second run cloned from second run if any,
    else from first run. Otherwise replace all runs with single run cloned from first run."""
    runs = p.runs
    if not runs:
        p.add_run(body_text)
        return

    if prefix_text is not None:
        # Keep first run's properties for prefix, second run (or first) for body.
        first = runs[0]
        first.text = prefix_text
        body_run_template = runs[1] if len(runs) > 1 else runs[0]
        # Clone the second run's rPr by adding a new run after the first
        # Easiest: set first run prefix, set second run body, delete the rest.
        if len(runs) == 1:
            new_run = p.add_run(body_text)
        else:
            runs[1].text = body_text
            # remove any later runs
            for r in runs[2:]:
                r._element.getparent().remove(r._element)
    else:
        # Single run replacement: keep first run's formatting, drop the rest.
        runs[0].text = body_text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)


def main():
    doc = Document(PATH)
    for idx, (prefix, body) in NEW.items():
        p = doc.paragraphs[idx]
        replace_paragraph_text(p, prefix, body)
    doc.save(PATH)
    print(f"Updated {len(NEW)} paragraphs.")


if __name__ == "__main__":
    main()
