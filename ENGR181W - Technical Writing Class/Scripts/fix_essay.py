import docx
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = docx.Document('Ethics_Impact_Essay.docx')

# Locate anchor: first body section "2.1.1 ..."
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('2.1.1'):
        anchor = p
        break
if anchor is None:
    raise SystemExit('Could not find 2.1.1 anchor')

heading_style = anchor.style  # 'Heading 1'
normal_style = None  # body paragraphs have no explicit style in this doc

def insert_before(anchor_p, text, style):
    new_p = OxmlElement('w:p')
    anchor_p._p.addprevious(new_p)
    new_para = Paragraph(new_p, anchor_p._parent)
    if style is not None:
        new_para.style = style
    new_para.add_run(text)
    return new_para

abstract_text = (
    "This essay examines the professional, ethical, environmental, and societal dimensions of the Modular Battery "
    "Management and Communication Platform, a 36 V lithium-ion battery management system (BMS) developed for the EE "
    "175AB senior design sequence at UC Riverside. The platform delivers up to 4 A through a 10-series Molicel P42A "
    "pack and is intended for personal electric mobility and embedded robotics applications. If commercialized, the "
    "system would be deployed in dense urban environments where lithium-ion fault events have caused fatalities and "
    "have driven new certification mandates in New York City and the European Union Battery Regulation 2023/1542. "
    "This paper identifies the public safety, privacy, and equity-of-access implications of such a deployment, "
    "describes how the design addresses each through a layered hardware and firmware safety architecture, reflects "
    "on the lessons the team learned about engineering honesty and documentation, and assesses the global, "
    "environmental, and societal impact of low-cost open-architecture battery management technology. The paper "
    "concludes that responsible commercialization requires treating public safety as the foundational design "
    "commitment rather than a constraint to be minimized."
)

keywords_text = (
    "battery management system, lithium-ion safety, micro-mobility, telemetry privacy, IEEE Code of Ethics"
)

intro_paragraphs = [
    (
        "The rapid global growth of lithium-ion energy storage has produced a corresponding increase in safety "
        "incidents, regulatory activity, and public concern surrounding battery-powered consumer products. Fires "
        "originating in electric scooters, e-bikes, and similar devices have caused injuries, fatalities, and "
        "structural damage in major cities, prompting new certification mandates and import restrictions [2]. At "
        "the same time, the electrification of transport remains one of the most cost-effective pathways for "
        "reducing greenhouse gas emissions in urbanizing regions [7]. Engineers who design battery management "
        "systems therefore operate at the intersection of two competing pressures: the need to deploy energy "
        "storage technology widely, and the obligation to ensure that each unit deployed will not endanger the "
        "public."
    ),
    (
        "The Modular Battery Management and Communication Platform is a senior design project that addresses this "
        "intersection directly. The system manages a 10-series lithium-ion pack with hardware-enforced fault "
        "protection, real-time telemetry, and a modular architecture intended for adaptation to electric scooters, "
        "e-bikes, and embedded robotics platforms. Although the prototype was developed in a controlled "
        "laboratory, a commercial version would operate in environments where the consequences of a single design "
        "failure can extend far beyond the individual user."
    ),
    (
        "This essay addresses four questions: (a) what ethical implications a commercial version of this system "
        "would carry; (b) how the present design responds to those implications; (c) what the design team learned "
        "about professional and ethical responsibility through the project; and (d) what global, economic, "
        "environmental, and societal impacts could follow from broader deployment of the platform. The remainder "
        "of the essay is organized into sections that examine deployment context, ethical implications, lessons "
        "learned, and broader impact, followed by concluding remarks and references."
    ),
]

insert_before(anchor, 'Abstract', heading_style)
insert_before(anchor, abstract_text, normal_style)
insert_before(anchor, 'Keywords', heading_style)
insert_before(anchor, keywords_text, normal_style)
insert_before(anchor, '1. Introduction', heading_style)
for para in intro_paragraphs:
    insert_before(anchor, para, normal_style)

# In-text citations: paragraph-level replacement (rebuild first run, clear rest)
replacements = [
    ('smaller college towns.', 'smaller college towns [1].'),
    ('California Consumer Privacy Act (CCPA)', 'California Consumer Privacy Act (CCPA) [3]'),
    ('European General Data Protection Regulation (GDPR).', 'European General Data Protection Regulation (GDPR) [4].'),
    ('EN 62133 for secondary lithium cells', 'EN 62133 for secondary lithium cells [5]'),
    ('mandating UL certification for lithium batteries sold or rented in the city.',
     'mandating UL certification for lithium batteries sold or rented in the city [2].'),
    ('safety, health, and welfare of the public.”', 'safety, health, and welfare of the public.” [6]'),
    ('claims or estimates based on available data.”', 'claims or estimates based on available data.” [6]'),
    ('International Energy Agency (IEA)', 'International Energy Agency (IEA) [7]'),
    ('Battery Regulation (EU) 2023/1542, which came into force in August 2023,',
     'Battery Regulation (EU) 2023/1542, which came into force in August 2023 [8],'),
    ('Consumer Product Safety Commission (CPSC) has issued guidance on lithium battery safety',
     'Consumer Product Safety Commission (CPSC) has issued guidance on lithium battery safety [9]'),
    ('IPC-2221 clearance and creepage requirements',
     'IPC-2221 clearance and creepage requirements [10]'),
    ('AEC-Q qualified components where possible',
     'AEC-Q qualified components where possible [11]'),
    ('Texas Instruments BQ76930 analog front end',
     'Texas Instruments BQ76930 analog front end [12]'),
]

applied = {old: False for old, _ in replacements}

for para in doc.paragraphs:
    full = para.text
    changed = False
    for old, new in replacements:
        if old in full:
            full = full.replace(old, new)
            applied[old] = True
            changed = True
    if changed:
        # Rewrite paragraph: keep first run formatting, replace text, clear other runs.
        if para.runs:
            para.runs[0].text = full
            for r in para.runs[1:]:
                r.text = ''
        else:
            para.add_run(full)

for old, ok in applied.items():
    if not ok:
        print(f'WARN: not applied: {old[:60]}')

# References
references = [
    '[1] Bird Global, Inc., "Sustainability Report," Miami, FL, USA, 2022. [Online]. Available: https://www.bird.co/sustainability/',
    '[2] New York City Fire Department, "FDNY Public Safety Bulletin: Lithium-Ion Battery Fires," New York, NY, USA, 2023. [Online]. Available: https://www.nyc.gov/site/fdny/',
    '[3] State of California, "California Consumer Privacy Act of 2018," Cal. Civ. Code, sec. 1798.100 et seq., 2018.',
    '[4] European Parliament and Council, "Regulation (EU) 2016/679 (General Data Protection Regulation)," Official Journal of the European Union, vol. L 119, pp. 1-88, May 2016.',
    '[5] International Electrotechnical Commission, "IEC 62133-2: Secondary cells and batteries containing alkaline or other non-acid electrolytes. Safety requirements for portable sealed secondary lithium cells," IEC, Geneva, Switzerland, 2017.',
    '[6] IEEE, "IEEE Code of Ethics," IEEE Policies, sec. 7.8, 2020. [Online]. Available: https://www.ieee.org/about/corporate/governance/p7-8.html',
    '[7] International Energy Agency, "Global EV Outlook 2024," IEA, Paris, France, 2024. [Online]. Available: https://www.iea.org/reports/global-ev-outlook-2024',
    '[8] European Parliament and Council, "Regulation (EU) 2023/1542 concerning batteries and waste batteries," Official Journal of the European Union, vol. L 191, pp. 1-117, Jul. 2023.',
    '[9] U.S. Consumer Product Safety Commission, "Letter to Manufacturers, Importers, Distributors, and Retailers of Lithium-Ion Batteries and Products Containing Such Batteries," CPSC, Bethesda, MD, USA, Dec. 2022.',
    '[10] IPC, "IPC-2221B: Generic Standard on Printed Board Design," IPC, Bannockburn, IL, USA, 2012.',
    '[11] Automotive Electronics Council, "AEC-Q100 Rev-H: Failure Mechanism Based Stress Test Qualification for Integrated Circuits," AEC, 2014.',
    '[12] Texas Instruments, "BQ76930 3-Series to 15-Series Cell Battery Monitor and Protector," Datasheet SLUSBK3, Dallas, TX, USA, 2018. [Online]. Available: https://www.ti.com/product/BQ76930',
]

ref_heading = doc.add_paragraph('References')
ref_heading.style = heading_style
for r in references:
    doc.add_paragraph(r)

# Final em-dash sweep just in case
for para in doc.paragraphs:
    for run in para.runs:
        if '—' in run.text:
            run.text = run.text.replace('—', ' - ')

doc.save('Ethics_Impact_Essay.docx')
print('Done.')
