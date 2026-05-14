"""
Rebuild Ethics_Impact_Essay_v4.docx with humanized prose, citation fixes
already applied, and IEEE two-column format. Uses the same layout helpers
as build_v3.py.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell

FONT  = 'Times New Roman'
COL_W = 3.525
GAP_TWP = 144


def _run(p, text, size=10, bold=False, italic=False, sc=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if sc:
        rPr = r._r.get_or_add_rPr()
        rPr.append(OxmlElement('w:smallCaps'))
    return r


def _fmt(p, align, sb, sa, li, ri, fi):
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    if li is not None: pf.left_indent = Inches(li)
    if ri is not None: pf.right_indent = Inches(ri)
    if fi is not None: pf.first_line_indent = Inches(fi)


class W:
    def __init__(self, container):
        self._c = container
        self._first = isinstance(container, _Cell)

    def _next(self):
        if self._first:
            self._first = False
            return self._c.paragraphs[0]
        return self._c.add_paragraph()

    def para(self, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=4, li=None, ri=None, fi=None):
        p = self._next()
        _fmt(p, align, sb, sa, li, ri, fi)
        return p

    def body(self, text):
        p = self.para(sb=0, sa=3, fi=0.18)
        _run(p, text)

    def heading(self, text):
        p = self.para(align=WD_ALIGN_PARAGRAPH.CENTER, sb=8, sa=4)
        _run(p, text, sc=True)

    def ref(self, num, text):
        p = self.para(sb=0, sa=2, li=0.22, fi=-0.22)
        _run(p, num + ' ', size=8)
        _run(p, text, size=8)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tb = OxmlElement('w:tblBorders')
    for s in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{s}')
        el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'auto')
        tb.append(el)
    tblPr.append(tb)
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            cb = OxmlElement('w:tcBorders')
            for s in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                el = OxmlElement(f'w:{s}')
                el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0')
                el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'auto')
                cb.append(el)
            tcPr.append(cb)


def _set_table_layout(table, col_widths_in):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    tblW = OxmlElement('w:tblW')
    total = sum(col_widths_in) * 1440
    tblW.set(qn('w:w'), str(int(total))); tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths_in:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w * 1440)))
        tblGrid.append(gc)
    tbl.insert(tbl.index(tblPr) + 1, tblGrid)


def _cell_width(cell, inches):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(inches * 1440))); tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _cell_margins(cell, left=0, right=0, top=0, bot=0):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    m = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bot), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val)); el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(0.625); sec.right_margin = Inches(0.625)
sp = sec._sectPr
for c in sp.findall(qn('w:cols')):
    sp.remove(c)

d = W(doc)

# Title
p = d.para(align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
_run(p, 'Understanding of Professional and Ethical Responsibility and Global,\n'
        'Economic, Environmental, and Societal Impact', size=24)

p = d.para(align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=6)
_run(p, 'Modular Battery Management and Communication Platform', size=16)

p = d.para(align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=2)
_run(p, 'Kaushik Vada, Connor Stewart, Nick Poon, Christian Kim', size=11)

p = d.para(align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=8)
_run(p, 'UC Riverside, EE 175AB, 2026', size=11, italic=True)

# Two-column body
table = doc.add_table(rows=1, cols=2, style='Table Grid')
_no_borders(table)
_set_table_layout(table, [COL_W, COL_W])
L, R = table.cell(0, 0), table.cell(0, 1)
_cell_width(L, COL_W); _cell_margins(L, left=0, right=GAP_TWP)
_cell_width(R, COL_W); _cell_margins(R, left=GAP_TWP, right=0)

lc = W(L); rc = W(R)

# Abstract
p = lc.para(sa=5)
_run(p, 'Abstract—', size=9, bold=True, italic=True)
_run(p,
    'A battery management system decides whether a lithium pack is safe, useful, '
    'or on fire, so the choices behind it are not only technical. This essay covers '
    'the professional, ethical, environmental, and societal questions raised by the '
    'Modular Battery Management and Communication Platform built for the EE 175AB '
    'senior design sequence at UC Riverside. The platform supplies 36 V at up to 4 A '
    'from a 10-series Molicel P42A pack and is aimed at electric scooters, e-bikes, '
    'skateboards, and embedded robotics. A commercial version would operate in dense '
    'urban environments where lithium-ion fires have already killed people and '
    'triggered new certification rules in New York City and the European Union. The '
    'design responds with layered hardware and firmware safeguards, honest performance '
    'reporting, and an open-source architecture meant to make safe battery management '
    'more accessible outside well-resourced markets.',
    size=9, italic=True)

p = lc.para(sa=10)
_run(p, 'Index Terms—', size=9, bold=True, italic=True)
_run(p,
    'battery management system, lithium-ion safety, micro-mobility, '
    'telemetry privacy, IEEE Code of Ethics',
    size=9, italic=True)

lc.heading('I.  Introduction')
lc.body(
    'Eighteen people died in New York City lithium-ion fires in 2023 [2]. The '
    'European Union has answered with certification mandates; most US cities are '
    'still writing theirs. At the same time, electrifying two-wheeled urban '
    'transport is one of the cheaper ways to cut transport emissions in growing '
    'cities [7]. Engineers building battery management systems sit between those '
    'two pressures, and the only way through is to take safety seriously from the '
    'first schematic.')
lc.body(
    'The Modular Battery Management and Communication Platform manages a 10-series '
    'lithium-ion pack with hardware-enforced fault protection and real-time '
    'telemetry. The same board can be adapted to electric scooters, e-bikes, or '
    'embedded robotics. The prototype was tested in a university lab. A commercial '
    'version would run in transit stations, on sidewalks, and in field environments, '
    'where a single design failure can hurt people who had nothing to do with the '
    'product.')
lc.body(
    'This essay covers four things: the ethical issues a commercial version raises, '
    'how the current design responds to them, what the team learned about '
    'professional responsibility while building it, and what wider deployment would '
    'mean.')

lc.heading('II.  Deployment Context')
lc.body(
    'The BMS outputs 36 V nominal at up to 4 A from a 10-series Molicel P42A pack. '
    'The prototype was tested in a controlled UC Riverside lab; commercial '
    'deployment is a different problem. Shared scooter and e-bike fleets now '
    'operate in cities across multiple continents [1], from tropical humidity to '
    'sub-freezing winters. The same board would also power embedded platforms like '
    'the NVIDIA Jetson Orin Nano in field robotics, where motor-drive EMI and '
    'sustained vibration are part of normal operation.')
lc.body(
    'End users range from riders who have never opened a battery enclosure to '
    'research engineers who watch the telemetry GUI without knowing much about what '
    'the hardware underneath is doing. The design team is responsible for both.')
lc.body(
    'A 151 Wh pack on a crowded sidewalk is not a thought experiment. When '
    'something fails at that energy density in a public space, the consequences '
    'land on bystanders who never chose to be near it. That shaped every protection '
    'decision in the design.')
lc.body(
    'Data carries its own obligations. The BMS currently sends telemetry over USB '
    'CDC in a closed lab; in a commercial fleet product, the equivalent wireless '
    'stream carries location and usage patterns that fall under CCPA [3] and GDPR '
    '[4]. Anonymization, encryption, and access control have to be planned before '
    'the data starts flowing, instead of bolted on afterward.')

lc.heading('III.  Ethical Implications of Commercialization')
lc.body(
    'The 18 lithium-ion fire deaths in New York City in 2023 led to local '
    'legislation requiring UL certification for any lithium battery sold or rented '
    'in the city [2]. That kind of response is what happens when safety-critical '
    'products reach the market without adequate protection across the pack’s '
    'full service life, instead of only at bench-nominal conditions.')
lc.body(
    'The design uses hardware-layered protection for that reason. The Texas '
    'Instruments BQ76930 analog front end [12] handles fault protection '
    'independently of the STM32 firmware. Short-circuit protection trips within 70 '
    'µs without any software involvement, in line with the short-circuit and '
    'overcharge safety framework IEC 62133-2 defines for portable lithium cells '
    '[5]. Firmware-only protection is exposed to software bugs and communication '
    'failures that have caused real incidents, so the critical faults are handled '
    'before firmware enters the picture.')
lc.body(
    'The switch from a MOSFET-based final breaker to a mechanical relay shows that '
    'priority in practice. Datasheet discrepancies showed the MOSFET design could '
    'not be relied on under certain fault conditions. The change was made under '
    'real schedule and budget pressure, in line with the IEEE Code of Ethics '
    'requirement to hold public safety paramount [6]. The budget was tight; the '
    'safety requirement was not negotiable.')
lc.body(
    'Open architecture also has ethical weight. The first-run prototype cost about '
    '$3,000 including development hardware. Separating the battery stack, firmware, '
    'and visualization layer makes the design easier to cost-reduce at volume and '
    'easier to adapt for different cell configurations. Hosting schematics and '
    'firmware on GitHub puts that knowledge in front of student teams and small '
    'operators who cannot afford commercial alternatives. Safe, understandable '
    'battery management is not a given in lower-resource contexts, and an open '
    'architecture chips away at that.')
lc.body(
    'Even in its current closed-lab form, the telemetry pipeline has a data-ethics '
    'obligation. Data minimization, purpose limitation, and user consent are '
    'architectural decisions, and retrofitting them onto an existing pipeline is '
    'harder than designing them in from the start. If the team ever builds a '
    'wireless fleet version, those choices should already be settled.')

# Right column
rc.heading('IV.  Professional and Ethical Responsibility: Lessons Learned')
rc.body(
    'The biggest design lesson was that fault handling cannot be retrofitted. It '
    'has to be designed in at the schematic stage. In hardware, the IEEE Code of '
    'Ethics requirement to hold public safety paramount [6] has a concrete meaning: '
    'every fault condition—overcurrent, overvoltage, undervoltage, short '
    'circuit, overtemperature—needed documented test cases at a 100% pass '
    'rate before the BMS connected to a live pack. The Cell Emulator PCB was built '
    'specifically so hardware fault testing did not expose the team to a live '
    'lithium pack. Protecting the engineers who validate the system is part of '
    'public safety, not a separate problem.')
rc.body(
    'Honest reporting was the harder discipline. Post-production routing errors '
    'in the E-Load PCB made the original 5 A target impossible. The temptation in '
    'that situation is to report numbers you cannot back up. The IEEE Code of '
    'Ethics is explicit that engineers must be honest in stating claims based on '
    'available data [6]. In a safety-critical system, overstating performance is a '
    'hazard, since users design their own systems around your specifications, and '
    'a wrong number in a datasheet can become a real failure in the field. The '
    'team documented the limitation, derated to 1 A, and described the fix. That '
    'is what honest engineering looks like under pressure.')
rc.body(
    'Documentation is engineering work. Version-controlled PCB files, a firmware '
    'repository with a real README, and operator installation guides are how '
    'knowledge about a safety-critical system survives the people who built it. '
    'In this kind of system, that continuity matters more than usual.')

rc.heading('V.  Global, Economic, Environmental, and Societal Impact')
rc.body(
    'The biggest global opportunity for this platform is in developing markets. '
    'Across sub-Saharan Africa, South and Southeast Asia, and Latin America, urban '
    'mobility runs on petrol-powered motorcycles and scooters, which are a major '
    'source of local air pollution and transport carbon. The IEA treats '
    'electrifying those fleets as one of the more accessible mitigation paths [7]. '
    'The bottleneck is not cells or motors. It is battery management that local '
    'manufacturers can source, repair, and actually understand. An open-architecture '
    'BMS built to IPC-2221 [10] with AEC-Q components [11] where applicable is '
    'useful in that context in a way a closed commercial unit is not.')
rc.body(
    'Cell balancing matters environmentally too. Keeping weaker cells from '
    'degrading ahead of the rest of the pack extends total pack life, which means '
    'fewer packs manufactured, less lithium and cobalt mined, and less battery '
    'waste at end of life. The benefit grows with adoption.')
rc.body(
    'Regulation is moving in one direction. The EU Battery Regulation (EU) '
    '2023/1542 [8] now requires BMS performance data, carbon footprint declarations, '
    'and end-of-life collection for EV batteries, including light electric '
    'vehicles. After a series of high-profile fires, the CPSC has urged '
    'manufacturers, importers, and retailers of lithium-ion batteries and '
    'e-mobility devices to certify compliance with consensus UL standards [9], and '
    'has since moved toward mandatory US standards. Early conformance to IPC-2221, '
    'hardware-independent fault protection, and clean signal integrity should make '
    'certification tractable without a structural redesign.')
rc.body(
    'In rural and peri-urban communities with limited grid access, reliable '
    'battery storage is not a convenience. It backs small businesses, medical '
    'refrigeration, and communications infrastructure. A BMS with clearly '
    'separated, documented layers is easier to adapt and easier to teach where '
    'engineering budgets are tight and staff learn by working directly with the '
    'hardware.')

rc.heading('VI.  Conclusion')
rc.body(
    'At a technical level, this is a 36 V lithium-ion safety and telemetry system. '
    'The harder part was the judgment calls: which fault conditions to protect '
    'against and how, how to handle a spec the hardware could not meet, and what '
    'to put in writing. Hardware-independent fault protection, honest performance '
    'reporting, open-source architecture, conformance to industry standards, and '
    'documented fault validation are not separate decisions. They are all what the '
    'IEEE Code of Ethics asks for in practice.')
rc.body(
    'The potential for affordable, verifiable battery management to matter '
    'globally is real. Whether it actually does depends on whether the engineers '
    'who build and sell these systems treat public safety as the first constraint '
    'in the design, instead of a compliance check at the end. The technical work '
    'was the straightforward part. The ethical commitments are what decide whether '
    'any of it matters.')

rc.heading('References')
REFS = [
    ('[1]',  'Bird Global, Inc., “Sustainability Report,” Miami, FL, USA, 2022. [Online]. Available: https://www.bird.co/sustainability/'),
    ('[2]',  'New York City Fire Department, “Lithium-Ion Battery Safety,” New York, NY, USA, 2023. [Online]. Available: https://www.nyc.gov/site/fdny/codes/reference/lithium-ion-battery-safety.page'),
    ('[3]',  'State of California, “California Consumer Privacy Act of 2018,” Cal. Civ. Code, sec. 1798.100 et seq., 2018.'),
    ('[4]',  'European Parliament and Council, “Regulation (EU) 2016/679 (General Data Protection Regulation),” Official Journal of the European Union, vol. L 119, pp. 1–88, May 2016.'),
    ('[5]',  'International Electrotechnical Commission, “IEC 62133-2: Secondary cells and batteries containing alkaline or other non-acid electrolytes. Safety requirements for portable sealed secondary lithium cells,” IEC, Geneva, Switzerland, 2017.'),
    ('[6]',  'IEEE, “IEEE Code of Ethics,” IEEE Policies, sec. 7.8, 2020. [Online]. Available: https://www.ieee.org/about/corporate/governance/p7-8'),
    ('[7]',  'International Energy Agency, “Global EV Outlook 2024,” IEA, Paris, France, 2024. [Online]. Available: https://www.iea.org/reports/global-ev-outlook-2024'),
    ('[8]',  'European Parliament and Council, “Regulation (EU) 2023/1542 concerning batteries and waste batteries,” Official Journal of the European Union, vol. L 191, pp. 1–117, Jul. 2023.'),
    ('[9]',  'U.S. Consumer Product Safety Commission, “Letter to Manufacturers, Importers, Distributors, and Retailers of Lithium-Ion Batteries and Products Containing Such Batteries,” CPSC, Bethesda, MD, USA, Dec. 2022.'),
    ('[10]', 'IPC, “IPC-2221B: Generic Standard on Printed Board Design,” IPC, Bannockburn, IL, USA, 2012.'),
    ('[11]', 'Automotive Electronics Council, “AEC-Q100 Rev-H: Failure Mechanism Based Stress Test Qualification for Integrated Circuits,” AEC, 2014.'),
    ('[12]', 'Texas Instruments, “BQ769x0 3-Series to 15-Series Cell Battery Monitor Family for Li-Ion and Phosphate Applications,” Datasheet SLUSBK2I, Dallas, TX, USA, Mar. 2022. [Online]. Available: https://www.ti.com/product/BQ76930'),
]
for num, text in REFS:
    rc.ref(num, text)

out = 'Ethics_Impact_Essay_v4.docx'
try:
    doc.save(out)
except PermissionError:
    out = 'Ethics_Impact_Essay_v4_humanized.docx'
    doc.save(out)
    print('(v4 was locked - wrote to', out, 'instead)')
print('Saved', out)
