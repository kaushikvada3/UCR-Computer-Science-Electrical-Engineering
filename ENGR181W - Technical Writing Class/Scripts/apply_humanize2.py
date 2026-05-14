"""Second-pass humanization: flatten residual parallelisms and vary conclusion rhythm."""
from docx import Document

PATH = r"C:\Users\chris_bptnrfq\Desktop\181W\ENGR181W - Technical Writing Class\Ethics_Impact_Essay.docx"

NEW = {}

# P21 - drop "rather than optional features" parallelism
NEW[21] = "The physical environment creates engineering problems that touch directly on safety. Lithium-ion cells are sensitive to temperature extremes. High ambient temperatures speed up capacity fade and raise the risk of thermal runaway; low temperatures cut available capacity and can pull cells below their safe minimum voltage if discharge continues. The BMS handles this with continuous NTC thermistor monitoring on every thermal channel, a PWM-controlled fan for active cooling, and a hardware-enforced breaker that disconnects the load above 60 °C. The team treated these as primary safety requirements from the start of the design."

# P26 - tighten "deliberate safety choice rather than an afterthought"; soften IEEE quote framing
NEW[26] = "The design uses a layered safety architecture for this reason. At the hardware level, the Texas Instruments BQ76930 analog front end [12] handles fault protection independently of the STM32 firmware. Short-circuit protection trips within 70 ms with no software involvement. This matters because firmware-dependent protection is exposed to software bugs, stack overflows, and communication failures, and all three have contributed to real-world battery incidents. The redundant mechanical relay that serves as the final breaker, picked to meet the current and voltage ratings of the 10s lithium pack, was a deliberate safety decision. The team replaced the original MOSFET-based breaker after datasheet discrepancies risked leaving the system exposed, and put in a validated mechanical relay to close the gap. The decision was made under time and budget pressure, and it is the kind of judgment the IEEE Code of Ethics asks for: hold public safety paramount even when it is inconvenient."

# P27 - reduce "real" repetition and soften
NEW[27] = "Fairness of access is a smaller but important concern. The BMS was built to replace expensive commercial units with a low-cost, open-architecture alternative. The first-run prototype cost about $2,966 including development hardware. The modular architecture separates the battery stack, the firmware, and the visualization layer, which makes the cost easier to bring down at volume and the system easier to adapt to different cell configurations. Reliable battery management that small-scale EV builders, student robotics teams, and low-income micro-mobility operators in developing markets can actually afford has a meaningful societal benefit. Hosting the firmware and GUI openly on GitHub follows the same engineering principle of sharing knowledge with the profession."

# P32 - drop the closing parallelism
NEW[32] = "The team also learned that documentation is part of professional responsibility. Keeping the PCB files in version control, a firmware repository with a real README, and release documentation that includes operator installation guides is not paperwork. It is how engineering knowledge gets preserved and reviewed, and it is what the next engineer needs to maintain or modify the system safely. Poor documentation of safety-critical systems has contributed to engineering disasters in fields ranging from industrial process control to medical devices. Documenting properly is part of the engineering work."

# P39 - rework conclusion: vary rhythm, drop "first commitment, not a constraint to push down"
NEW[39] = "Technically, this is a 36 V lithium-ion safety and telemetry system. The harder part of the project was ethical. Rechargeable energy storage is a domain where a single design failure shows up as fires, injuries, and lost public trust. The choices the team made (hardware-independent fault protection, performance numbers reported as they actually measured, an open-source architecture, conformance to industry standards, and validated behavior under each fault condition) line up with what the IEEE Code of Ethics asks of practicing engineers. The platform's global, environmental, and societal potential is real. Whether that potential matters in practice depends on whether the engineers who build and sell it put public safety first."


def replace_paragraph_text(p, body_text):
    runs = p.runs
    if not runs:
        p.add_run(body_text)
        return
    runs[0].text = body_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def main():
    doc = Document(PATH)
    for idx, body in NEW.items():
        replace_paragraph_text(doc.paragraphs[idx], body)
    doc.save(PATH)
    print(f"Updated {len(NEW)} paragraphs in second pass.")


if __name__ == "__main__":
    main()
